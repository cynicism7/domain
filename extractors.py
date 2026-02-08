# -*- coding: utf-8 -*-
"""从 PDF、DOCX、TXT 中提取「标题、摘要、正文前两页」，供领域识别使用。"""

import re
from pathlib import Path
from typing import Tuple


# 正文前两页约等于的字符数（中英文混合）
CHARS_PER_PAGE = 1200


def _find_abstract_and_body(full_text: str, abstract_max: int = 1500, body_max: int = 2400) -> Tuple[str, str]:
    """从全文里解析摘要段和正文前两页。"""
    text = (full_text or "").strip()
    abstract = ""
    body = ""
    # 摘要起始关键词（中英文）
    abstract_starts = re.compile(
        r"(?:^|\n)\s*(?:Abstract|ABSTRACT|摘要|【摘要】)\s*[：:\s]*\n?",
        re.IGNORECASE
    )
    # 摘要结束 / 正文开始关键词
    body_starts = re.compile(
        r"(?:Introduction|INTRODUCTION|1\.\s+Introduction|引言|前言|Keywords|Key words|索引词|I\.\s+)",
        re.IGNORECASE
    )
    match_start = abstract_starts.search(text)
    if match_start:
        after_abstract_label = text[match_start.end():]
        match_end = body_starts.search(after_abstract_label)
        if match_end:
            abstract = after_abstract_label[: match_end.start()].strip()[:abstract_max]
            body = after_abstract_label[match_end.start() :].strip()[:body_max]
        else:
            # 没有明确正文起始，摘要取到一定长度，正文用摘要后的内容
            abstract = after_abstract_label[:abstract_max].strip()
            body = after_abstract_label[abstract_max : abstract_max + body_max].strip()
    else:
        # 没有找到摘要标题，前一段当摘要，后面当正文前两页
        abstract = text[:abstract_max].strip()
        body = text[abstract_max : abstract_max + body_max].strip()
    return abstract, body


def extract_txt(path: str, max_chars: int = 5000) -> str:
    """从纯文本文件读取内容。"""
    path = Path(path)
    if not path.exists():
        return ""
    try:
        for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                text = path.read_text(encoding=enc)
                return text[:max_chars] if max_chars else text
            except UnicodeDecodeError:
                continue
    except Exception:
        pass
    return ""


def extract_pdf(path: str, max_chars: int = 5000) -> str:
    """从 PDF 提取前几页文本。"""
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    path = Path(path)
    if not path.exists():
        return ""
    try:
        reader = PdfReader(path)
        parts = []
        n = 0
        for page in reader.pages:
            if n >= 5:  # 最多 5 页
                break
            t = page.extract_text()
            if t:
                parts.append(t)
                n += 1
        text = "\n".join(parts)
        return text[:max_chars] if max_chars else text
    except Exception:
        return ""


def extract_docx(path: str, max_chars: int = 5000) -> str:
    """从 DOCX 提取段落文本。"""
    try:
        from docx import Document
    except ImportError:
        return ""
    path = Path(path)
    if not path.exists():
        return ""
    try:
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(parts)
        return text[:max_chars] if max_chars else text
    except Exception:
        return ""


def extract_text(file_path: str, max_chars: int = 5000) -> str:
    """根据扩展名选择提取器并返回文本（兼容旧接口）。"""
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return extract_txt(file_path, max_chars)
    if suffix == ".pdf":
        return extract_pdf(file_path, max_chars)
    if suffix in (".docx", ".doc"):
        return extract_docx(file_path, max_chars)
    return ""


def extract_title_abstract_body(
    file_path: str,
    abstract_max: int = 1500,
    body_pages_chars: int = None,
) -> Tuple[str, str, str]:
    """
    提取「标题、摘要、正文前两页」三部分，用于领域判断。
    返回 (title, abstract, body_first_2_pages)。
    """
    body_max = body_pages_chars or (CHARS_PER_PAGE * 2)
    path = Path(file_path)
    name = path.name
    suffix = path.suffix.lower()

    title = name
    full_text = ""

    if suffix == ".txt":
        full_text = extract_txt(file_path, max_chars=abstract_max + body_max + 2000)
        lines = full_text.splitlines()
        if lines:
            first_line = lines[0].strip()
            if len(first_line) < 200 and first_line:
                title = first_line
    elif suffix == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            meta = reader.metadata
            if meta and getattr(meta, "title", None) and str(meta.title).strip():
                title = str(meta.title).strip()
            pages_text = []
            for i, page in enumerate(reader.pages):
                if i >= 5:
                    break
                t = page.extract_text()
                if t:
                    pages_text.append(t)
            full_text = "\n".join(pages_text)
        except Exception:
            full_text = ""
    elif suffix in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(path)
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            if paras:
                first = paras[0].strip()
                if len(first) < 200 and first:
                    title = first
            full_text = "\n".join(paras)
        except Exception:
            full_text = ""
    else:
        full_text = extract_text(file_path, max_chars=abstract_max + body_max + 2000)

    abstract, body = _find_abstract_and_body(full_text, abstract_max=abstract_max, body_max=body_max)
    return title, abstract, body
